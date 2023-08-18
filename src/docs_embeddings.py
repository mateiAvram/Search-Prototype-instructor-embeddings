docs_path = 'FRDs/'

from docx import Document
import re
import os

from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.vectorstores import Chroma

def find_docx_files(docs_path):
    docx_files = []
    
    for filename in os.listdir(docs_path):
        if filename.endswith(".docx"):
            docx_files.append(filename)
    
    return docx_files

docx_files = []
if os.path.isdir(docs_path):
    docx_files = find_docx_files(docs_path)

docs = {}
doc_id = 0
for file in docx_files:

    # file_path = docs_path + docx_files[0]
    file_path = docs_path + file
    doc = Document(docx=file_path)

    texts = []
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading'):
            continue
        text = p.text.strip()
        if text:
            texts.append(text)

    chunks = []
    if texts != []:

        i = 0
        while i < len(texts):
            flag = True
            count = i
            chunk_length = 0
            chunk = texts[count]
            while flag:
                # Quick fix
                if count + 1 == len(texts):
                    flag = False
                    continue

                if len(chunk + texts[count + 1]) > 511:
                    flag = False
                    continue
                chunk = chunk + '\n' + texts[count + 1]
                count += 1
            
            chunks.append(chunk)
            i += 1        

    # # print(f'file name: {file}')
    # print('chunks:')
    # for chunk in chunks: 
    #     print(chunk)
    #     print()

    # print('---------')
    # print()
    # print()

    docs[file_path] = (str(doc_id), chunks)
    doc_id += 1

ef = HuggingFaceInstructEmbeddings()
vectorstore = Chroma(persist_directory='embeddings', collection_name='tasklight', embedding_function=ef)

for key in docs.keys():
    t = docs[key]
    path = key.split('/')
    file_name = path[-1]
    
    document_id = t[0]
    document_name = file_name[:-5]
    texts = t[1]

    metadata = {'document_id': document_id, 'file_path': key}

    if texts == []:
        # Title of file is used for embeddings
        title_metadata = metadata.copy()
        title_metadata['type'] = 'title'

        vectorstore.add_texts(texts=[document_name], metadatas=[title_metadata])
        print(f'Added file_name embeddings for: {document_name}')
    else:
        # Texts from the file are used for embeddings
        texts_metadatas = []
        for text in texts:
            text_metadata = metadata.copy()
            text_metadata['type'] = 'text'
            texts_metadatas.append(text_metadata)
        
        vectorstore.add_texts(texts=texts, metadatas=texts_metadatas)
        print(f'Added embeddings for texts inside document with file_name: {document_name}')

##########################################################################################################################

# class TreeNode:
#     def __init__(self, data, style=None):
#         self.data = data
#         self.style = style
#         self.children = []

#     def add_child(self, child_node):
#         self.children.append(child_node)

#     def __repr__(self, level=0):
#         ret = "\t" * level + f"{repr(self.data)}\n"
#         for child in self.children:
#             ret += child.__repr__(level + 1)
#         return ret

# def parse_document(doc_path):
#     doc = Document(doc_path)
#     root = TreeNode("Document Root")

#     stack = [root]

#     for paragraph in doc.paragraphs:
#         if paragraph.style.name.startswith('Heading'):
#             level = int(paragraph.style.name.split()[-1])
#             text = paragraph.text.strip()
#             style = f"heading {level}"
#             current_node = TreeNode(text, style)
#             while len(stack) > level:
#                 stack.pop()
#             stack[-1].add_child(current_node)
#             stack.append(current_node)
#         else:
#             text = paragraph.text.strip()
#             if text:
#                 current_node = TreeNode(text, "paragraph")
#                 stack[-1].add_child(current_node)

#     return root

# doc_path = 'FRDs/Taubate.docx'
# tree = parse_document(doc_path=doc_path)

# print(tree)
